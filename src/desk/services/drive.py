"""Google Drive API wrapper."""

import csv
import io
from pathlib import Path

from docx import Document
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from openpyxl import load_workbook


class DriveClient:
    """Client for Google Drive API operations."""

    def __init__(self, credentials: Credentials):
        self.service = build("drive", "v3", credentials=credentials)

    # ── Shared Drive helpers ─────────────────────────────────────────
    # Centralize supportsAllDrives / includeItemsFromAllDrives so every
    # API call works with Shared Drives automatically.  See ADR-015.

    def _files_list(
        self,
        drive_id: str | None = None,
        my_drive: bool = False,
        **kwargs,
    ):
        """Wrap files().list() with Shared Drive parameters."""
        kwargs["supportsAllDrives"] = True
        kwargs["includeItemsFromAllDrives"] = True
        if drive_id:
            kwargs["corpora"] = "drive"
            kwargs["driveId"] = drive_id
        elif my_drive:
            kwargs["corpora"] = "user"
        else:
            kwargs["corpora"] = "allDrives"
        return self.service.files().list(**kwargs)

    def _files_get(self, **kwargs):
        """Wrap files().get() with Shared Drive support."""
        kwargs["supportsAllDrives"] = True
        return self.service.files().get(**kwargs)

    def _files_get_media(self, **kwargs):
        """Wrap files().get_media() with Shared Drive support."""
        kwargs["supportsAllDrives"] = True
        return self.service.files().get_media(**kwargs)

    def _files_create(self, **kwargs):
        """Wrap files().create() with Shared Drive support."""
        kwargs["supportsAllDrives"] = True
        return self.service.files().create(**kwargs)

    def _files_update(self, **kwargs):
        """Wrap files().update() with Shared Drive support."""
        kwargs["supportsAllDrives"] = True
        return self.service.files().update(**kwargs)

    def _files_copy(self, **kwargs):
        """Wrap files().copy() with Shared Drive support."""
        kwargs["supportsAllDrives"] = True
        return self.service.files().copy(**kwargs)

    def _files_export(self, **kwargs):
        """Wrap files().export() — export does not accept supportsAllDrives."""
        return self.service.files().export(**kwargs)

    def _files_export_media(self, **kwargs):
        """Wrap files().export_media() — export does not accept supportsAllDrives."""
        return self.service.files().export_media(**kwargs)

    def _permissions_create(self, **kwargs):
        """Wrap permissions().create() with Shared Drive support."""
        kwargs["supportsAllDrives"] = True
        return self.service.permissions().create(**kwargs)

    def _permissions_list(self, **kwargs):
        """Wrap permissions().list() with Shared Drive support."""
        kwargs["supportsAllDrives"] = True
        return self.service.permissions().list(**kwargs)

    def _permissions_delete(self, **kwargs):
        """Wrap permissions().delete() with Shared Drive support."""
        kwargs["supportsAllDrives"] = True
        return self.service.permissions().delete(**kwargs)

    # ── Public methods ───────────────────────────────────────────────

    def search(
        self,
        query: str,
        max_results: int = 20,
        page_token: str | None = None,
        drive_id: str | None = None,
        my_drive: bool = False,
    ) -> dict:
        """Search for files matching query.

        Args:
            query: Drive search query (e.g., "name contains 'report'")
            max_results: Maximum number of results
            page_token: Token for fetching next page of results
            drive_id: Scope search to a specific Shared Drive
            my_drive: Scope search to My Drive only (faster)

        Returns:
            Dict with 'files' list and 'nextPageToken' (if more results exist)
        """
        try:
            request_kwargs = {
                "q": query,
                "pageSize": max_results,
                "fields": (
                    "nextPageToken, files(id, name, mimeType, modifiedTime, viewedByMeTime, size, owners, webViewLink)"
                ),
                "orderBy": "modifiedTime desc",
            }
            if page_token:
                request_kwargs["pageToken"] = page_token

            results = self._files_list(
                drive_id=drive_id, my_drive=my_drive, **request_kwargs
            ).execute()

            result = {"files": results.get("files", [])}
            if results.get("nextPageToken"):
                result["nextPageToken"] = results["nextPageToken"]
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    # Mime types for uploaded Office files
    _DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    _XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    def read(self, file_id: str) -> str:
        """Read file content as text (exports Google Docs/Sheets as plain text).

        Google Workspace files are exported via the API. Uploaded Office files
        (.docx, .xlsx) are downloaded and converted locally using python-docx
        and openpyxl.

        Args:
            file_id: The file ID

        Returns:
            File content as string
        """
        try:
            # First get file metadata to determine type
            meta = self._files_get(fileId=file_id, fields="mimeType, name").execute()
            mime = meta["mimeType"]

            # Google Workspace files need export
            if mime == "application/vnd.google-apps.document":
                return self._export(file_id, "text/plain")
            elif mime == "application/vnd.google-apps.spreadsheet":
                return self._export(file_id, "text/csv")
            elif mime == "application/vnd.google-apps.presentation":
                return self._export(file_id, "text/plain")
            else:
                # Check file size before downloading (50MB limit)
                size_meta = self._files_get(fileId=file_id, fields="size").execute()
                size = int(size_meta.get("size", 0))
                if size > 50 * 1024 * 1024:
                    raise RuntimeError(
                        f"File is {size // (1024 * 1024)}MB — too large to read as text. "
                        "Use 'desk drive download' for large files."
                    )
                content = self._files_get_media(fileId=file_id).execute()
                if not isinstance(content, bytes):
                    content = str(content).encode("utf-8")

                # Route Office files to local converters
                if mime == self._DOCX_MIME:
                    return _read_docx(content)
                elif mime == self._XLSX_MIME:
                    return _read_xlsx(content)

                return content.decode("utf-8", errors="replace")
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def info(self, file_id: str) -> dict:
        """Get detailed file metadata.

        Args:
            file_id: The file ID

        Returns:
            File metadata dict
        """
        try:
            return (
                self._files_get(
                    fileId=file_id,
                    fields="id, name, mimeType, modifiedTime, viewedByMeTime, createdTime, size, owners, "
                    "parents, webViewLink, description, starred, driveId",
                )
                .execute()
            )
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def recent(
        self,
        max_results: int = 20,
        page_token: str | None = None,
        drive_id: str | None = None,
        my_drive: bool = False,
    ) -> dict:
        """List recently modified files.

        Args:
            max_results: Maximum number of results
            page_token: Token for fetching next page of results
            drive_id: Scope to a specific Shared Drive
            my_drive: Scope to My Drive only (faster)

        Returns:
            Dict with 'files' list and 'nextPageToken' (if more results exist)
        """
        try:
            request_kwargs = {
                "pageSize": max_results,
                "fields": "nextPageToken, files(id, name, mimeType, modifiedTime, viewedByMeTime, webViewLink)",
                "orderBy": "modifiedTime desc",
            }
            if page_token:
                request_kwargs["pageToken"] = page_token

            results = self._files_list(
                drive_id=drive_id, my_drive=my_drive, **request_kwargs
            ).execute()

            result = {"files": results.get("files", [])}
            if results.get("nextPageToken"):
                result["nextPageToken"] = results["nextPageToken"]
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def upload(self, local_path: str, folder_id: str | None = None) -> dict:
        """Upload a local file to Drive.

        Args:
            local_path: Path to local file
            folder_id: Optional parent folder ID

        Returns:
            File metadata dict
        """
        path = Path(local_path)
        if not path.exists():
            raise RuntimeError(f"File not found: {local_path}")

        file_metadata = {"name": path.name}
        if folder_id:
            file_metadata["parents"] = [folder_id]

        try:
            media = MediaFileUpload(str(path))
            result = self._files_create(
                body=file_metadata,
                media_body=media,
                fields="id, name, webViewLink",
            ).execute()
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def download(self, file_id: str, dest_path: str) -> str:
        """Download a file to local filesystem.

        Args:
            file_id: The file ID
            dest_path: Local directory or file path to save to

        Returns:
            Path to downloaded file
        """
        try:
            meta = self._files_get(fileId=file_id, fields="name, mimeType").execute()
            name = meta["name"]
            mime = meta["mimeType"]

            dest = Path(dest_path)
            # Sanitize filename to prevent path traversal
            safe_name = Path(name).name
            if dest.is_dir():
                dest = dest / safe_name

            # Google Workspace files need export
            export_map = {
                "application/vnd.google-apps.document": ("application/pdf", ".pdf"),
                "application/vnd.google-apps.spreadsheet": ("text/csv", ".csv"),
                "application/vnd.google-apps.presentation": ("application/pdf", ".pdf"),
            }

            if mime in export_map:
                export_mime, ext = export_map[mime]
                if not dest.suffix:
                    dest = dest.with_suffix(ext)
                request = self._files_export_media(fileId=file_id, mimeType=export_mime)
            else:
                request = self._files_get_media(fileId=file_id)

            with io.FileIO(str(dest), "wb") as fh:
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()

            return str(dest)
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def mkdir(self, name: str, parent_id: str | None = None) -> dict:
        """Create a folder in Drive.

        Args:
            name: Folder name
            parent_id: Optional parent folder ID

        Returns:
            Folder metadata dict
        """
        file_metadata = {
            "name": name,
            "mimeType": "application/vnd.google-apps.folder",
        }
        if parent_id:
            file_metadata["parents"] = [parent_id]

        try:
            result = (
                self._files_create(body=file_metadata, fields="id, name, webViewLink")
                .execute()
            )
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def move(self, file_id: str, folder_id: str) -> dict:
        """Move a file to a different folder.

        Args:
            file_id: The file to move
            folder_id: Destination folder ID

        Returns:
            Updated file metadata
        """
        try:
            # Get current parents to remove
            file = self._files_get(fileId=file_id, fields="parents").execute()
            previous_parents = ",".join(file.get("parents", []))

            result = (
                self._files_update(
                    fileId=file_id,
                    addParents=folder_id,
                    removeParents=previous_parents,
                    fields="id, name, parents, webViewLink",
                )
                .execute()
            )
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def trash(self, file_id: str) -> dict:
        """Move a file to trash.

        Args:
            file_id: The file ID

        Returns:
            Updated file metadata
        """
        try:
            result = self._files_update(
                fileId=file_id,
                body={"trashed": True},
                fields="id, name, trashed",
            ).execute()
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def untrash(self, file_id: str) -> dict:
        """Restore a file from trash.

        Args:
            file_id: The file ID

        Returns:
            Updated file metadata
        """
        try:
            result = self._files_update(
                fileId=file_id,
                body={"trashed": False},
                fields="id, name, trashed",
            ).execute()
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def share(self, file_id: str, email: str, role: str = "writer") -> dict:
        """Share a file with someone.

        Args:
            file_id: The file ID
            email: Email address to share with
            role: Permission role (reader, commenter, writer)

        Returns:
            Permission dict
        """
        try:
            permission = (
                self._permissions_create(
                    fileId=file_id,
                    body={"type": "user", "role": role, "emailAddress": email},
                    sendNotificationEmail=True,
                    fields="id, role, emailAddress",
                )
                .execute()
            )
            return permission
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def star(self, file_id: str, starred: bool = True) -> dict:
        """Star or unstar a file.

        Args:
            file_id: The file ID
            starred: True to star, False to unstar

        Returns:
            Updated file metadata
        """
        try:
            result = (
                self._files_update(
                    fileId=file_id,
                    body={"starred": starred},
                    fields="id, name, starred",
                )
                .execute()
            )
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def copy(
        self, file_id: str, name: str | None = None, folder_id: str | None = None
    ) -> dict:
        """Copy a file in Drive.

        Args:
            file_id: The file ID to copy
            name: Optional new name for the copy
            folder_id: Optional destination folder ID

        Returns:
            New file metadata dict
        """
        try:
            body = {}
            if name:
                body["name"] = name
            if folder_id:
                body["parents"] = [folder_id]

            result = (
                self._files_copy(
                    fileId=file_id,
                    body=body if body else None,
                    fields="id, name, webViewLink, parents",
                )
                .execute()
            )
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def list_permissions(self, file_id: str) -> list[dict]:
        """List all permissions on a file.

        Args:
            file_id: The file ID

        Returns:
            List of permission dicts
        """
        try:
            results = (
                self._permissions_list(
                    fileId=file_id,
                    fields="permissions(id, type, role, emailAddress, displayName, domain)",
                )
                .execute()
            )
            return results.get("permissions", [])
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def unshare(self, file_id: str, email: str) -> bool:
        """Remove a user's access to a file.

        Args:
            file_id: The file ID
            email: Email address to remove

        Returns:
            True if permission was removed

        Raises:
            ValueError: If user not found in permissions
        """
        # Find permission by email
        permissions = self.list_permissions(file_id)
        permission_id = None
        for perm in permissions:
            if perm.get("emailAddress", "").lower() == email.lower():
                permission_id = perm["id"]
                break

        if not permission_id:
            raise ValueError(f"User '{email}' not found in file permissions")

        try:
            self._permissions_delete(
                fileId=file_id, permissionId=permission_id
            ).execute()
            return True
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def transfer_ownership(self, file_id: str, email: str) -> dict:
        """Transfer ownership of a file to another user.

        Args:
            file_id: The file ID
            email: Email address of new owner

        Returns:
            Updated permission dict
        """
        try:
            # Create new owner permission
            result = (
                self._permissions_create(
                    fileId=file_id,
                    transferOwnership=True,
                    body={"type": "user", "role": "owner", "emailAddress": email},
                    fields="id, role, emailAddress",
                )
                .execute()
            )
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def list_comments(self, file_id: str, include_resolved: bool = False) -> list[dict]:
        """List comments on a file.

        Args:
            file_id: The file ID
            include_resolved: Include resolved comments

        Returns:
            List of comment dicts
        """
        try:
            results = (
                self.service.comments()
                .list(
                    fileId=file_id,
                    fields="comments(id, author, content, createdTime, resolved, quotedFileContent, replies)",
                )
                .execute()
            )
            comments = results.get("comments", [])

            # Filter resolved if needed
            if not include_resolved:
                comments = [c for c in comments if not c.get("resolved")]

            # Parse into clean format
            return [
                {
                    "id": c.get("id", ""),
                    "author": c.get("author", {}).get("displayName", ""),
                    "authorEmail": c.get("author", {}).get("emailAddress", ""),
                    "content": c.get("content", ""),
                    "createdTime": c.get("createdTime", ""),
                    "resolved": c.get("resolved", False),
                    "anchor": c.get("quotedFileContent", {}).get("value", ""),
                    "replyCount": len(c.get("replies", [])),
                }
                for c in comments
            ]
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def add_comment(self, file_id: str, content: str) -> dict:
        """Add a comment to a file.

        Args:
            file_id: The file ID
            content: Comment text

        Returns:
            Created comment dict
        """
        try:
            result = (
                self.service.comments()
                .create(
                    fileId=file_id,
                    body={"content": content},
                    fields="id, author, content, createdTime",
                )
                .execute()
            )
            return {
                "id": result.get("id", ""),
                "author": result.get("author", {}).get("displayName", ""),
                "content": result.get("content", ""),
                "createdTime": result.get("createdTime", ""),
            }
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def resolve_comment(self, file_id: str, comment_id: str, resolved: bool = True) -> dict:
        """Resolve or reopen a comment.

        Args:
            file_id: The file ID
            comment_id: The comment ID
            resolved: True to resolve, False to reopen

        Returns:
            Updated comment dict
        """
        try:
            # Get the current comment content (required for update)
            current = (
                self.service.comments()
                .get(fileId=file_id, commentId=comment_id, fields="content")
                .execute()
            )

            result = (
                self.service.comments()
                .update(
                    fileId=file_id,
                    commentId=comment_id,
                    body={"content": current.get("content", "")},
                    fields="id, content, resolved",
                )
                .execute()
            )

            # Resolve/reopen by updating via replies if needed
            # The comments API resolves via the resolved field
            if resolved:
                # Create a reply that resolves
                self.service.replies().create(
                    fileId=file_id,
                    commentId=comment_id,
                    body={"content": "", "action": "resolve"},
                    fields="id",
                ).execute()
            else:
                # Reopen by creating a reply with reopen action
                self.service.replies().create(
                    fileId=file_id,
                    commentId=comment_id,
                    body={"content": "", "action": "reopen"},
                    fields="id",
                ).execute()

            return {"id": comment_id, "resolved": resolved}
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def reply_comment(self, file_id: str, comment_id: str, content: str) -> dict:
        """Reply to a comment.

        Args:
            file_id: The file ID
            comment_id: The comment ID to reply to
            content: Reply text

        Returns:
            Created reply dict
        """
        try:
            result = (
                self.service.replies()
                .create(
                    fileId=file_id,
                    commentId=comment_id,
                    body={"content": content},
                    fields="id, author, content, createdTime",
                )
                .execute()
            )
            return {
                "id": result.get("id", ""),
                "author": result.get("author", {}).get("displayName", ""),
                "content": result.get("content", ""),
                "createdTime": result.get("createdTime", ""),
            }
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    # Friendly name → MIME type prefix mapping for --type filter
    _TYPE_MIME_MAP = {
        "document": "application/vnd.google-apps.document",
        "spreadsheet": "application/vnd.google-apps.spreadsheet",
        "presentation": "application/vnd.google-apps.presentation",
        "pdf": "application/pdf",
        "folder": "application/vnd.google-apps.folder",
        "image": "image/",
        "video": "video/",
        "audio": "audio/",
    }

    def list_folder(
        self,
        folder_id: str,
        max_results: int = 100,
        page_token: str | None = None,
        file_type: str | None = None,
        drive_id: str | None = None,
        my_drive: bool = False,
    ) -> dict:
        """List files in a folder with external pagination.

        Filters out sub-folders and shortcuts at the query level (unless
        file_type explicitly requests folders).

        Args:
            folder_id: The folder ID
            max_results: Maximum number of files to return (default 100)
            page_token: Token for fetching next page of results
            file_type: Optional friendly type filter (e.g. 'document', 'spreadsheet', 'pdf')
            drive_id: Scope to a specific Shared Drive
            my_drive: Scope to My Drive only (faster)

        Returns:
            Dict with 'files' list and optional 'nextPageToken'
        """
        query = f"'{folder_id}' in parents and trashed=false"

        if file_type:
            mime = self._TYPE_MIME_MAP.get(file_type.lower())
            if mime:
                if mime.endswith("/"):
                    # Prefix match (e.g. image/*)
                    query += f" and mimeType contains '{mime}'"
                else:
                    query += f" and mimeType = '{mime}'"
            else:
                # Treat as raw MIME type
                query += f" and mimeType = '{file_type}'"
        else:
            # Default: exclude folders and shortcuts
            query += (
                " and mimeType != 'application/vnd.google-apps.folder'"
                " and mimeType != 'application/vnd.google-apps.shortcut'"
            )

        try:
            request_kwargs = {
                "q": query,
                "pageSize": min(max_results, 100),
                "fields": (
                    "nextPageToken, files(id, name, mimeType, modifiedTime, viewedByMeTime, size)"
                ),
                "orderBy": "modifiedTime desc",
            }
            if page_token:
                request_kwargs["pageToken"] = page_token

            results = self._files_list(
                drive_id=drive_id, my_drive=my_drive, **request_kwargs
            ).execute()

            result = {"files": results.get("files", [])}
            if results.get("nextPageToken"):
                result["nextPageToken"] = results["nextPageToken"]
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def list_drives(
        self, max_results: int = 100, page_token: str | None = None
    ) -> dict:
        """List available Shared Drives (single page).

        Args:
            max_results: Page size (max 100)
            page_token: Token for fetching next page of results

        Returns:
            Dict with 'drives' list and optional 'nextPageToken'
        """
        if max_results <= 0:
            return {"drives": []}
        try:
            kwargs: dict = {
                "pageSize": min(max_results, 100),
                "fields": "nextPageToken, drives(id, name, createdTime)",
            }
            if page_token:
                kwargs["pageToken"] = page_token
            results = self.service.drives().list(**kwargs).execute()

            result: dict = {"drives": results.get("drives", [])}
            if results.get("nextPageToken"):
                result["nextPageToken"] = results["nextPageToken"]
            return result
        except HttpError as error:
            raise RuntimeError(f"Drive API error: {error}")

    def _export(self, file_id: str, mime_type: str) -> str:
        """Export a Google Workspace file."""
        content = self._files_export(fileId=file_id, mimeType=mime_type).execute()
        if isinstance(content, bytes):
            return content.decode("utf-8")
        return str(content)


def _read_docx(content: bytes) -> str:
    """Extract text from a .docx file, including paragraphs and tables.

    Args:
        content: Raw .docx file bytes

    Returns:
        Extracted text as string
    """
    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        raise RuntimeError(
            f"Could not read .docx file (may be encrypted or corrupted): {e}. "
            "Use 'desk drive download' to save the raw file."
        )

    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("\t".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n".join(parts)


def _read_xlsx(content: bytes) -> str:
    """Extract data from an .xlsx file as CSV text.

    Multiple sheets are separated with markers.

    Args:
        content: Raw .xlsx file bytes

    Returns:
        CSV text with sheet separators
    """
    try:
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as e:
        raise RuntimeError(
            f"Could not read .xlsx file (may be encrypted or corrupted): {e}. "
            "Use 'desk drive download' to save the raw file."
        )

    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if len(wb.sheetnames) > 1:
            parts.append(f"--- sheet: {sheet_name} ---")

        output = io.StringIO()
        writer = csv.writer(output)
        for row in ws.iter_rows(values_only=True):
            writer.writerow([str(cell) if cell is not None else "" for cell in row])
        parts.append(output.getvalue().rstrip())

    wb.close()
    return "\n".join(parts)
