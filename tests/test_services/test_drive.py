"""Tests for Drive service client."""

import io
import pytest
from unittest.mock import MagicMock, patch

from googleapiclient.errors import HttpError


class TestDriveClientInit:
    """Tests for DriveClient initialization."""

    def test_creates_service_with_credentials(self, mock_credentials):
        """Should create Drive service with provided credentials."""
        with patch("desk.services.drive.build") as mock_build:
            mock_build.return_value = MagicMock()
            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)

            mock_build.assert_called_once_with("drive", "v3", credentials=mock_credentials)


class TestDriveSearch:
    """Tests for DriveClient.search method."""

    def test_search_returns_files(self, mock_credentials):
        """Should return list of files matching query."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {
                "files": [
                    {"id": "file1", "name": "Report.docx"},
                    {"id": "file2", "name": "Data.xlsx"},
                ]
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.search("name contains 'Report'")

            assert "files" in result
            assert len(result["files"]) == 2

    def test_search_with_max_results(self, mock_credentials):
        """Should pass max_results as pageSize to API."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.search("name contains 'test'", max_results=10)

            call_kwargs = files_mock.list.call_args[1]
            assert call_kwargs["pageSize"] == 10

    def test_search_api_error_raises_runtime_error(self, mock_credentials):
        """Should raise RuntimeError on API error."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            http_error = HttpError(
                resp=MagicMock(status=400),
                content=b'{"error": {"message": "Invalid query"}}'
            )
            files_mock.list.return_value.execute.side_effect = http_error

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            with pytest.raises(RuntimeError, match="Drive API error"):
                client.search("invalid:query")


class TestDriveInfo:
    """Tests for DriveClient.info method."""

    def test_info_returns_file_metadata(self, mock_credentials):
        """Should return detailed file metadata."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.get.return_value.execute.return_value = {
                "id": "file123",
                "name": "Document.docx",
                "mimeType": "application/vnd.google-apps.document",
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.info("file123")

            assert result["id"] == "file123"
            assert result["name"] == "Document.docx"


class TestDriveRecent:
    """Tests for DriveClient.recent method."""

    def test_recent_returns_files(self, mock_credentials):
        """Should return list of recently modified files."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {
                "files": [
                    {"id": "file1", "name": "Recent1.docx"},
                    {"id": "file2", "name": "Recent2.docx"},
                ]
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.recent()

            assert "files" in result
            assert len(result["files"]) == 2


class TestDriveTrash:
    """Tests for DriveClient trash operations."""

    def test_trash_updates_file(self, mock_credentials):
        """Should update file with trashed=True."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.update.return_value.execute.return_value = {"id": "file123"}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.trash("file123")

            files_mock.update.assert_called_once()
            call_kwargs = files_mock.update.call_args[1]
            assert call_kwargs["body"]["trashed"] is True

    def test_untrash_updates_file(self, mock_credentials):
        """Should update file with trashed=False."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.update.return_value.execute.return_value = {"id": "file123"}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.untrash("file123")

            call_kwargs = files_mock.update.call_args[1]
            assert call_kwargs["body"]["trashed"] is False


class TestDriveStar:
    """Tests for DriveClient star operations."""

    def test_star_updates_file(self, mock_credentials):
        """Should update file with starred=True."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.update.return_value.execute.return_value = {"id": "file123"}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.star("file123")

            call_kwargs = files_mock.update.call_args[1]
            assert call_kwargs["body"]["starred"] is True


class TestDriveShare:
    """Tests for DriveClient.share method."""

    def test_share_creates_permission(self, mock_credentials):
        """Should create permission for email."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            permissions_mock = mock_service.permissions.return_value
            permissions_mock.create.return_value.execute.return_value = {
                "id": "perm123",
                "role": "reader",
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.share("file123", "user@example.com", role="reader")

            permissions_mock.create.assert_called_once()
            call_kwargs = permissions_mock.create.call_args[1]
            assert call_kwargs["body"]["emailAddress"] == "user@example.com"
            assert call_kwargs["body"]["role"] == "reader"


class TestDriveListFolder:
    """Tests for DriveClient.list_folder method."""

    def test_list_folder_returns_files(self, mock_credentials):
        """Should return dict with files list."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {
                "files": [
                    {"id": "f1", "name": "Doc.docx", "mimeType": "application/vnd.google-apps.document"},
                    {"id": "f2", "name": "Sheet.xlsx", "mimeType": "application/vnd.google-apps.spreadsheet"},
                ]
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.list_folder("folder123")

            assert "files" in result
            assert len(result["files"]) == 2

    def test_list_folder_with_pagination(self, mock_credentials):
        """Should pass page_token and return nextPageToken."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {
                "files": [{"id": "f1", "name": "Doc.docx"}],
                "nextPageToken": "token_page2",
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.list_folder("folder123", page_token="token_page1")

            call_kwargs = files_mock.list.call_args[1]
            assert call_kwargs["pageToken"] == "token_page1"
            assert result["nextPageToken"] == "token_page2"

    def test_list_folder_with_max_results(self, mock_credentials):
        """Should pass max_results as pageSize (capped at 100)."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.list_folder("folder123", max_results=10)

            call_kwargs = files_mock.list.call_args[1]
            assert call_kwargs["pageSize"] == 10

    def test_list_folder_caps_page_size_at_100(self, mock_credentials):
        """Should cap pageSize at 100 even if max_results is higher."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.list_folder("folder123", max_results=500)

            call_kwargs = files_mock.list.call_args[1]
            assert call_kwargs["pageSize"] == 100

    def test_list_folder_with_file_type_filter(self, mock_credentials):
        """Should add MIME type filter to query for known types."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.list_folder("folder123", file_type="document")

            call_kwargs = files_mock.list.call_args[1]
            assert "application/vnd.google-apps.document" in call_kwargs["q"]
            # Should not exclude folders when type is specified
            assert "folder" not in call_kwargs["q"] or "google-apps.document" in call_kwargs["q"]

    def test_list_folder_with_raw_mime_type(self, mock_credentials):
        """Should accept raw MIME type strings when not a known friendly name."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.list_folder("folder123", file_type="application/pdf")

            call_kwargs = files_mock.list.call_args[1]
            assert "application/pdf" in call_kwargs["q"]

    def test_list_folder_empty(self, mock_credentials):
        """Should return empty files list for empty folder."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.list_folder("empty_folder")

            assert result["files"] == []
            assert "nextPageToken" not in result

    def test_list_folder_excludes_folders_by_default(self, mock_credentials):
        """Should exclude folders and shortcuts when no file_type specified."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.list_folder("folder123")

            call_kwargs = files_mock.list.call_args[1]
            assert "vnd.google-apps.folder" in call_kwargs["q"]
            assert "vnd.google-apps.shortcut" in call_kwargs["q"]

    def test_list_folder_api_error_raises_runtime_error(self, mock_credentials):
        """Should raise RuntimeError on API error."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            http_error = HttpError(
                resp=MagicMock(status=404),
                content=b'{"error": {"message": "Folder not found"}}'
            )
            files_mock.list.return_value.execute.side_effect = http_error

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            with pytest.raises(RuntimeError, match="Drive API error"):
                client.list_folder("bad_folder")


class TestReadDocx:
    """Tests for _read_docx helper."""

    def test_read_docx_extracts_paragraphs(self):
        """Should extract paragraph text from a .docx file."""
        from docx import Document as DocxDocument

        # Create a real docx in memory
        doc = DocxDocument()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        from desk.services.drive import _read_docx

        result = _read_docx(content)
        assert "Hello world" in result
        assert "Second paragraph" in result

    def test_read_docx_extracts_tables(self):
        """Should extract table content from a .docx file."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        from desk.services.drive import _read_docx

        result = _read_docx(content)
        assert "A1" in result
        assert "B2" in result

    def test_read_docx_error_on_corrupt_file(self):
        """Should raise RuntimeError on corrupt .docx content."""
        from desk.services.drive import _read_docx

        with pytest.raises(RuntimeError, match="Could not read .docx file"):
            _read_docx(b"not a docx file")


class TestReadXlsx:
    """Tests for _read_xlsx helper."""

    def test_read_xlsx_single_sheet(self):
        """Should extract data from single-sheet .xlsx file as CSV."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["Name", "Value"])
        ws.append(["Alice", 42])
        buf = io.BytesIO()
        wb.save(buf)
        content = buf.getvalue()

        from desk.services.drive import _read_xlsx

        result = _read_xlsx(content)
        assert "Name" in result
        assert "Alice" in result
        assert "42" in result
        # Single sheet should NOT have sheet separator
        assert "--- sheet:" not in result

    def test_read_xlsx_multi_sheet(self):
        """Should include sheet separators for multi-sheet .xlsx files."""
        from openpyxl import Workbook

        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sales"
        ws1.append(["Q1", 100])

        ws2 = wb.create_sheet("Costs")
        ws2.append(["Q1", 50])

        buf = io.BytesIO()
        wb.save(buf)
        content = buf.getvalue()

        from desk.services.drive import _read_xlsx

        result = _read_xlsx(content)
        assert "--- sheet: Sales ---" in result
        assert "--- sheet: Costs ---" in result
        assert "100" in result
        assert "50" in result

    def test_read_xlsx_error_on_corrupt_file(self):
        """Should raise RuntimeError on corrupt .xlsx content."""
        from desk.services.drive import _read_xlsx

        with pytest.raises(RuntimeError, match="Could not read .xlsx file"):
            _read_xlsx(b"not an xlsx file")


class TestDriveReadMimeRouting:
    """Tests for MIME type routing in DriveClient.read()."""

    def test_read_routes_docx_to_local_converter(self, mock_credentials):
        """Should route uploaded .docx files to _read_docx."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value

            # First call: get metadata
            docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            files_mock.get.return_value.execute.side_effect = [
                {"mimeType": docx_mime, "name": "Report.docx"},
                {"size": "1000"},
            ]

            # Create real docx bytes
            from docx import Document as DocxDocument

            doc = DocxDocument()
            doc.add_paragraph("Test content")
            buf = io.BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()

            files_mock.get_media.return_value.execute.return_value = docx_bytes

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.read("file123")

            assert "Test content" in result

    def test_read_routes_xlsx_to_local_converter(self, mock_credentials):
        """Should route uploaded .xlsx files to _read_xlsx."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value

            xlsx_mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            files_mock.get.return_value.execute.side_effect = [
                {"mimeType": xlsx_mime, "name": "Data.xlsx"},
                {"size": "2000"},
            ]

            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            ws.append(["Col1", "Col2"])
            ws.append(["val1", "val2"])
            buf = io.BytesIO()
            wb.save(buf)
            xlsx_bytes = buf.getvalue()

            files_mock.get_media.return_value.execute.return_value = xlsx_bytes

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.read("file123")

            assert "Col1" in result
            assert "val2" in result

    def test_read_exports_google_doc_as_text(self, mock_credentials):
        """Should export Google Docs as plain text."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.get.return_value.execute.return_value = {
                "mimeType": "application/vnd.google-apps.document",
                "name": "Native Doc",
            }
            files_mock.export.return_value.execute.return_value = b"Exported text"

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.read("gdoc123")

            assert result == "Exported text"
            files_mock.export.assert_called_once_with(
                fileId="gdoc123", mimeType="text/plain"
            )


class TestSharedDriveSupport:
    """Tests for Shared Drive (supportsAllDrives) support across methods."""

    def test_search_includes_shared_drive_params(self, mock_credentials):
        """Search should include supportsAllDrives, includeItemsFromAllDrives, corpora."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.search("test")

            call_kwargs = files_mock.list.call_args[1]
            assert call_kwargs["supportsAllDrives"] is True
            assert call_kwargs["includeItemsFromAllDrives"] is True
            assert call_kwargs["corpora"] == "allDrives"

    def test_search_with_drive_id(self, mock_credentials):
        """Search with drive_id should set corpora=drive and driveId."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.search("test", drive_id="0ABcDeFgHiJ")

            call_kwargs = files_mock.list.call_args[1]
            assert call_kwargs["corpora"] == "drive"
            assert call_kwargs["driveId"] == "0ABcDeFgHiJ"
            assert call_kwargs["supportsAllDrives"] is True

    def test_search_with_my_drive(self, mock_credentials):
        """Search with my_drive=True should set corpora=user."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.search("test", my_drive=True)

            call_kwargs = files_mock.list.call_args[1]
            assert call_kwargs["corpora"] == "user"
            assert "driveId" not in call_kwargs

    def test_info_includes_supports_all_drives(self, mock_credentials):
        """Info should pass supportsAllDrives=True."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.get.return_value.execute.return_value = {
                "id": "file123", "name": "Test"
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.info("file123")

            call_kwargs = files_mock.get.call_args[1]
            assert call_kwargs["supportsAllDrives"] is True

    def test_trash_includes_supports_all_drives(self, mock_credentials):
        """Trash should pass supportsAllDrives=True."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.update.return_value.execute.return_value = {"id": "file123"}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.trash("file123")

            call_kwargs = files_mock.update.call_args[1]
            assert call_kwargs["supportsAllDrives"] is True

    def test_share_includes_supports_all_drives(self, mock_credentials):
        """Share should pass supportsAllDrives=True on permissions.create."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            permissions_mock = mock_service.permissions.return_value
            permissions_mock.create.return_value.execute.return_value = {
                "id": "perm123", "role": "writer"
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.share("file123", "user@example.com")

            call_kwargs = permissions_mock.create.call_args[1]
            assert call_kwargs["supportsAllDrives"] is True

    def test_upload_includes_supports_all_drives(self, mock_credentials, tmp_path):
        """Upload should pass supportsAllDrives=True."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.create.return_value.execute.return_value = {
                "id": "new123", "name": "test.txt"
            }

            test_file = tmp_path / "test.txt"
            test_file.write_text("hello")

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.upload(str(test_file))

            call_kwargs = files_mock.create.call_args[1]
            assert call_kwargs["supportsAllDrives"] is True

    def test_copy_includes_supports_all_drives(self, mock_credentials):
        """Copy should pass supportsAllDrives=True."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.copy.return_value.execute.return_value = {
                "id": "copy123", "name": "Copy"
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.copy("file123")

            call_kwargs = files_mock.copy.call_args[1]
            assert call_kwargs["supportsAllDrives"] is True

    def test_list_folder_includes_shared_drive_params(self, mock_credentials):
        """list_folder should include Shared Drive params."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.list_folder("folder123")

            call_kwargs = files_mock.list.call_args[1]
            assert call_kwargs["supportsAllDrives"] is True
            assert call_kwargs["includeItemsFromAllDrives"] is True
            assert call_kwargs["corpora"] == "allDrives"

    def test_recent_includes_shared_drive_params(self, mock_credentials):
        """recent should include Shared Drive params."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            files_mock = mock_service.files.return_value
            files_mock.list.return_value.execute.return_value = {"files": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            client.recent()

            call_kwargs = files_mock.list.call_args[1]
            assert call_kwargs["supportsAllDrives"] is True
            assert call_kwargs["includeItemsFromAllDrives"] is True
            assert call_kwargs["corpora"] == "allDrives"


class TestListDrives:
    """Tests for DriveClient.list_drives method."""

    def test_list_drives_returns_drives(self, mock_credentials):
        """Should return dict with drives list."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            drives_mock = mock_service.drives.return_value
            drives_mock.list.return_value.execute.return_value = {
                "drives": [
                    {"id": "drive1", "name": "Engineering"},
                    {"id": "drive2", "name": "Design"},
                ]
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.list_drives()

            assert "drives" in result
            assert len(result["drives"]) == 2
            assert result["drives"][0]["name"] == "Engineering"

    def test_list_drives_empty(self, mock_credentials):
        """Should return empty drives list when no Shared Drives."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            drives_mock = mock_service.drives.return_value
            drives_mock.list.return_value.execute.return_value = {"drives": []}

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.list_drives()

            assert result["drives"] == []

    def test_list_drives_returns_next_page_token(self, mock_credentials):
        """Should return nextPageToken when the API returns one."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            drives_mock = mock_service.drives.return_value
            drives_mock.list.return_value.execute.return_value = {
                "drives": [
                    {"id": "drive1", "name": "Engineering"},
                ],
                "nextPageToken": "token_page2",
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.list_drives()

            assert result["nextPageToken"] == "token_page2"

    def test_list_drives_passes_page_token_to_api(self, mock_credentials):
        """Should pass page_token parameter to the API call."""
        with patch("desk.services.drive.build") as mock_build:
            mock_service = MagicMock()
            mock_build.return_value = mock_service

            drives_mock = mock_service.drives.return_value
            drives_mock.list.return_value.execute.return_value = {
                "drives": [
                    {"id": "drive3", "name": "Marketing"},
                ]
            }

            from desk.services.drive import DriveClient

            client = DriveClient(mock_credentials)
            result = client.list_drives(page_token="token_page1")

            call_kwargs = drives_mock.list.call_args[1]
            assert call_kwargs["pageToken"] == "token_page1"
            assert len(result["drives"]) == 1
